"""
report_exporter.py
Builds a downloadable Word (.docx) report containing:
- title and dataset overview
- the AI-generated insights text (parsed into sections)
- embedded charts

Uses python-docx since this module is part of the application
code itself (not a one-off document deliverable).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

BLUE = RGBColor(0x1A, 0x52, 0x76)
GRAY = RGBColor(0x55, 0x55, 0x55)


def _add_heading(doc, text, size=16, color=BLUE):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _parse_sections(insights_text: str) -> dict:
    """
    Splits the LLM's plain-text report into sections based on the
    known headers requested in the system prompt.
    """
    headers = [
        "Executive Summary",
        "Key Trends & Patterns",
        "Notable Anomalies / Outliers",
        "Actionable Recommendations",
    ]
    sections = {h: "" for h in headers}
    current = None

    for line in insights_text.splitlines():
        stripped = line.strip()
        matched = next((h for h in headers if stripped.startswith(h)), None)
        if matched:
            current = matched
            continue
        if current:
            sections[current] += line + "\n"

    return sections


def build_report(
    output_path: str,
    dataset_name: str,
    profile: dict,
    insights_text: str,
    chart_paths: list,
):
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("InsightGPT — AI Data Analysis Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"Dataset: {dataset_name}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = GRAY

    doc.add_paragraph()

    # Dataset overview
    _add_heading(doc, "Dataset Overview")
    overview = doc.add_paragraph()
    overview.add_run(
        f"Rows: {profile['shape']['rows']}    Columns: {profile['shape']['columns']}\n"
    )
    overview.add_run(f"Numeric columns: {', '.join(profile['numeric_columns']) or 'None'}\n")
    overview.add_run(f"Categorical columns: {', '.join(profile['categorical_columns']) or 'None'}")

    doc.add_paragraph()

    # AI Insights sections
    sections = _parse_sections(insights_text)
    for header, content in sections.items():
        _add_heading(doc, header, size=14)
        for line in content.strip().splitlines():
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style="List Bullet" if line.startswith(("-", "•")) else None)
            p.add_run(line.lstrip("-•").strip())
        doc.add_paragraph()

    # Charts
    if chart_paths:
        _add_heading(doc, "Visualizations")
        for chart in chart_paths:
            doc.add_paragraph(chart["title"]).runs[0].bold = True
            doc.add_picture(chart["path"], width=Inches(5.5))
            doc.add_paragraph()

    doc.save(output_path)
    return output_path
